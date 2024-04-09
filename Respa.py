import gradio as gr
import os
import pytesseract
from pdf2image import convert_from_path
from docx import Document
import spacy
import re
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from collections import Counter

# Load English language model for SpaCy
nlp = spacy.load("en_core_web_sm")

# Path to Tesseract executable (Update this with your Tesseract installation path)
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract'
poppler_path = r"C:\\Program Files (x86)\\poppler-23.07.0\\Library\\bin"
os.environ["PATH"] += os.pathsep + poppler_path

# Function to extract text from different file types
def extract_text_from_file(file):
    file_extension = file.name.split('.')[-1].lower()
    if file_extension == 'pdf':
        return extract_text_from_pdf(file)
    elif file_extension == 'docx':
        return extract_text_from_docx(file)
    elif file_extension == 'doc':
        return extract_text_from_doc(file)
    elif file_extension == 'txt':
        return extract_text_from_txt(file)
    else:
        return None

# Function to extract text from image-based PDF files using Tesseract OCR
def extract_text_from_pdf(pdf_file):
    try:
        # Convert the PDF pages to images
        images = convert_from_path(pdf_file.name, dpi=200)

        pdf_text = ""
        for page_number, image in enumerate(images):
            # Extract text using Tesseract OCR
            page_text = pytesseract.image_to_string(image, lang='eng')
            pdf_text += page_text + '\n'

        return pdf_text
    except Exception as e:
        return f"Error extracting text from PDF: {str(e)}"

# Function to extract text from a DOCX file
def extract_text_from_docx(file):
    try:
        doc = Document(file)
        doc_text = ""
        for para in doc.paragraphs:
            doc_text += para.text + '\n'
        return doc_text
    except Exception as e:
        return f"Error extracting text from DOCX: {str(e)}"

# Function to extract text from a DOC file
def extract_text_from_doc(file):
    try:
        doc = Document(file)
        doc_text = ""
        for para in doc.paragraphs:
            doc_text += para.text + '\n'
        return doc_text
    except Exception as e:
        return f"Error extracting text from DOC: {str(e)}"

# Function to extract text from a TXT file
def extract_text_from_txt(file):
    try:
        with open(file.name, 'r', encoding='utf-8') as txt_file:
            return txt_file.read()
    except Exception as e:
        return f"Error extracting text from TXT: {str(e)}"

# Function to extract entities from text
def extract_entities(text):
    doc = nlp(text)
    entities = [(ent.text, ent.label_) for ent in doc.ents]
    return entities

# Function to extract keywords from text
def extract_keywords(text):
    stop_words = set(stopwords.words('english'))
    words = word_tokenize(text)
    filtered_words = [word.lower() for word in words if word.isalnum() and word.lower() not in stop_words]
    word_freq = Counter(filtered_words)
    keywords = [word for word, freq in word_freq.items() if freq > 1]  # Adjust frequency threshold as needed
    return keywords

# Function to extract contact information from text
def extract_contact_info(text):
    contact_info = {
        'Location': "",
        'Name': "",
        'Phone': "",
        'Email': "",
        'LinkedIn': "",
        'GitHub': ""
    }

    # Extracting contact location
    location_regex = r'\b(?:location|address|city|state|country)\b\s*:\s*(.*)'
    location_match = re.search(location_regex, text, re.IGNORECASE)
    if location_match:
        contact_info['Location'] = location_match.group(1).strip()

    # Extracting contact name
    name_regex = r'\b(?:name|full name)\b\s*:\s*(.*)'
    name_match = re.search(name_regex, text, re.IGNORECASE)
    if name_match:
        contact_info['Name'] = name_match.group(1).strip()

    # Extracting contact phone
    phone_regex = r'\b(?:phone|mobile|telephone|contact)\b\s*:\s*(.*)'
    phone_match = re.search(phone_regex, text, re.IGNORECASE)
    if phone_match:
        contact_info['Phone'] = phone_match.group(1).strip()

    # Extracting contact email
    email_regex = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    email_match = re.search(email_regex, text)
    if email_match:
        contact_info['Email'] = email_match.group(0).strip()

    # Extracting LinkedIn URL
    linkedin_regex = r'\b(?:linkedin|linkedin profile|linkedin url)\b\s*:\s*(.*)'
    linkedin_match = re.search(linkedin_regex, text, re.IGNORECASE)
    if linkedin_match:
        contact_info['LinkedIn'] = linkedin_match.group(1).strip()

    # Extracting GitHub URL
    github_regex = r'\b(?:github|github profile|github url)\b\s*:\s*(.*)'
    github_match = re.search(github_regex, text, re.IGNORECASE)
    if github_match:
        contact_info['GitHub'] = github_match.group(1).strip()

    return contact_info

# Function to extract experience years from text
def extract_experience_years(text):
    years = re.findall(r'\b\d{4}\b', text)
    if years:
        start_year = min(map(int, years))
        end_year = max(map(int, years))
        return end_year - start_year
    else:
        return 0

# Function to map entities to keys
def map_entities_to_keys(entities):
    mapped_entities = {
        'Name': [],
        'ExperienceCompany': [],
        'Location': []
    }
    for text, label in entities:
        if label == 'PERSON':
            mapped_entities['Name'].append(text)
        elif label == 'ORG':
            mapped_entities['ExperienceCompany'].append(text)
        elif label == 'GPE':
            mapped_entities['Location'].append(text)
    return mapped_entities

# Function to map keywords to keys
def map_keywords_to_keys(keywords):
    mapped_keywords = {
        'Qualification': [],
        'Experience': False
    }
    if 'experience' in keywords:
        mapped_keywords['Experience'] = True
    return mapped_keywords

def extract_phone_numbers(text):
    phone_regex = r'(?:(?:\+?(\d{1,3}))?[\s-]?\(?\d{3}\)?[\s-]?\d{3}[\s-]?\d{4})'
    phone_match = re.search(phone_regex, text)
    phone = phone_match.group(0).strip() if phone_match else ""
    return phone

# Function to extract education details from text
def extract_education_details(text):
    education_degree_keywords = ['degree', 'qualification', 'education']
    education_institute_keywords = ['school', 'university', 'college', 'institute']
    education_start_end_time_regex = r'\b(\d{4})\b'

    doc = nlp(text)

    education_degrees = []
    education_institutes = []
    education_start_end_times = []

    for ent in doc.ents:
        if any(keyword in ent.text.lower() for keyword in education_degree_keywords):
            education_degrees.append(ent.text)
        if any(keyword in ent.text.lower() for keyword in education_institute_keywords):
            education_institutes.append(ent.text)
        if re.match(education_start_end_time_regex, ent.text):
            education_start_end_times.append(ent.text)

    return {
        'EducationDegree': education_degrees,
        'EducationInstitute': education_institutes,
        'EducationStartandEndTime': education_start_end_times
    }

# Function to extract experience details from text
def extract_experience_details(text):
    experience_company_keywords = ['company', 'organization', 'employer']
    experience_designation_keywords = ['position', 'title', 'role']
    experience_start_end_time_regex = r'\b(\d{4})\b-\b(\d{4})\b'

    doc = nlp(text)

    experience_companies = []
    experience_designations = []
    experience_start_end_times = []

    for ent in doc.ents:
        if any(keyword in ent.text.lower() for keyword in experience_company_keywords):
            experience_companies.append(ent.text)
        if any(keyword in ent.text.lower() for keyword in experience_designation_keywords):
            experience_designations.append(ent.text)
        if re.match(experience_start_end_time_regex, ent.text):
            experience_start_end_times.append(ent.text)

    return {
        'ExperienceCompany': experience_companies,
        'ExperienceDesignation': experience_designations,
        'ExperienceStartandEndTime': experience_start_end_times
    }

# Define function for Gradio interface
def parse_resume(file):
    resume_text = extract_text_from_file(file)

    education_details = extract_education_details(resume_text)
    experience_details = extract_experience_details(resume_text)
    contact_info = extract_contact_info(resume_text)
    entities = extract_entities(resume_text)
    keywords = extract_keywords(resume_text)

    mapped_entities = map_entities_to_keys(entities)
    mapped_keywords = map_keywords_to_keys(keywords)

    # Format education start and end times
    formatted_education_start_end_times = [f'"{time}"' for time in education_details['EducationStartandEndTime']]

    output = ""
    output += f"Name: {', '.join(mapped_entities['Name'])}\n"
    output += f"EducationDegree: {', '.join(education_details['EducationDegree'])}\n"
    output += f"EducationInstitute: {', '.join(education_details['EducationInstitute'])}\n"
    output += f"EducationStartandEndTime: [{', '.join(formatted_education_start_end_times)}]\n"
    output += f"Email: {contact_info['Email']}\n"
    output += f"ExperienceCompany: {', '.join(experience_details['ExperienceCompany'])}\n"
    output += f"ExperienceDesignation: {', '.join(experience_details['ExperienceDesignation'])}\n"
    output += f"ExperienceStartandEndTime: {', '.join(experience_details['ExperienceStartandEndTime'])}\n"
    output += f"Location: {', '.join(mapped_entities['Location'])}\n"
    output += f"Phone: {extract_phone_numbers(resume_text)}\n"

    # Add LinkedIn and GitHub URLs if available
    if contact_info['LinkedIn']:
        output += f"LinkedIn: {contact_info['LinkedIn']}\n"
    if contact_info['GitHub']:
        output += f"GitHub: {contact_info['GitHub']}\n"

    return output.strip()

# Define Gradio interface
iface = gr.Interface(parse_resume,
                     inputs="file",
                     outputs="text",
                     title="📝 RESUME PARSER",
                     description="Upload a resume file to parse its details.")

iface.launch()
